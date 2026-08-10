from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document
from workbench.parsers.docx_parser import DocumentParseError, parse_docx, write_outline_artifact


def build_outline(path: Path) -> None:
    document = Document()
    document.add_heading("一、专业概览", level=1)
    document.add_paragraph("")
    document.add_paragraph("这是正文段落。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "合并表头"
    table.cell(1, 0).text = "课程"
    table.cell(1, 1).text = "高等数学"
    document.add_heading("（二）培养方向", level=2)
    document.save(path)


def test_parse_docx_preserves_body_order_levels_tables_and_source_refs(tmp_path: Path) -> None:
    path = tmp_path / "大纲.docx"
    build_outline(path)

    outline = parse_docx(path)

    assert [(block.kind, block.order, block.level, block.text) for block in outline.blocks] == [
        ("heading", 1, 1, "一、专业概览"),
        ("paragraph", 2, None, "这是正文段落。"),
        ("table", 3, None, "合并表头 | 课程 | 高等数学"),
        ("heading", 4, 2, "（二）培养方向"),
    ]
    assert outline.blocks[2].table_cells == [["合并表头", "合并表头"], ["课程", "高等数学"]]
    assert [block.source_ref for block in outline.blocks] == [
        "paragraph:1",
        "paragraph:3",
        "table:1",
        "paragraph:4",
    ]


def test_parse_docx_writes_deterministic_artifact_with_cache_key(tmp_path: Path) -> None:
    path = tmp_path / "大纲.docx"
    build_outline(path)
    target = tmp_path / "03_文字识别" / "大纲结构.json"

    result = write_outline_artifact(path, target)
    payload = json.loads(target.read_text(encoding="utf-8"))

    assert payload["source_sha256"] == result.source_sha256
    assert payload["cache_key"] == f"docx-v1:{result.source_sha256}"
    assert payload["document"]["blocks"][0]["text"] == "一、专业概览"


def test_parse_docx_rejects_corrupt_package(tmp_path: Path) -> None:
    path = tmp_path / "损坏.docx"
    path.write_bytes(b"PK\x03\x04not-a-docx")

    with pytest.raises(DocumentParseError, match="无法解析"):
        parse_docx(path)
