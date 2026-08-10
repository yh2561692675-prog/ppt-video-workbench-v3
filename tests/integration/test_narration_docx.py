from __future__ import annotations

from datetime import UTC, datetime, timedelta
from pathlib import Path
from uuid import UUID, uuid4
from zipfile import ZipFile

import fitz
from docx import Document
from PIL import Image
from workbench.domain.confirmation import Confirmation
from workbench.domain.enums import NodeStatus
from workbench.domain.extraction import PageExtraction
from workbench.domain.models import NarrationRecord, PageRecord, ProjectManifest
from workbench.exports.narration_docx import export_narration_docx
from workbench.renderers.office_renderer import render_office_to_pdf


def _confirmed_project(project_dir: Path, *, long_page: bool = False) -> ProjectManifest:
    for folder in ("02_页面预览", "04_旁白", "08_输出"):
        (project_dir / folder).mkdir(parents=True, exist_ok=True)
    now = datetime.now(UTC)
    project_id = uuid4()
    pages: list[PageRecord] = []
    extractions: list[PageExtraction] = []
    confirmations: list[Confirmation] = []
    for order in range(1, 9):
        page_id = UUID(int=order)
        revision_id = UUID(int=100 + order)
        text = (
            "这是用于验证长段落不会截断的旁白。" * 90
            if long_page and order == 4
            else f"这是第{order}页的确认旁白，内容严格来自导入材料。"
        )
        pages.append(
            PageRecord(
                id=page_id,
                order=order,
                title=f"第{order}页主题",
                status=NodeStatus.COMPLETED,
                narration=NarrationRecord(
                    id=revision_id,
                    revision_id=revision_id,
                    text=text,
                    status=NodeStatus.COMPLETED,
                    confirmed_revision_id=revision_id,
                    author="规划师",
                    version=1,
                    updated_at=now,
                ),
            )
        )
        preview = None
        if order != 5:
            preview = project_dir / "02_页面预览" / f"page-{order:04d}.png"
            Image.new("RGB", (640, 360), (12 * order, 42, 72)).save(preview)
        extractions.append(
            PageExtraction(
                id=page_id,
                order=order,
                title=f"第{order}页主题",
                text=f"第{order}页课件文字",
                preview_path=(preview.relative_to(project_dir) if preview else None),
                extraction_method="image",
                source_ref=f"page:{order}",
            )
        )
        confirmations.append(
            Confirmation(
                id=uuid4(),
                page_id=page_id,
                revision_id=revision_id,
                actor="规划师",
                confirmed_at=now + timedelta(seconds=order),
            )
        )
    return ProjectManifest(
        id=project_id,
        name="八页专业介绍",
        project_dir=project_dir.name,
        created_at=now,
        updated_at=now,
        pages=list(reversed(pages)),
        page_extractions=list(reversed(extractions)),
        narration_confirmations=confirmations,
    )


def test_export_orders_pages_uses_chinese_font_and_handles_missing_thumbnail(
    tmp_path: Path,
) -> None:
    project_dir = tmp_path / "八页专业介绍_20260803_1200"
    manifest = _confirmed_project(project_dir)

    result = export_narration_docx(manifest, project_dir)

    assert result == project_dir / "04_旁白" / "旁白确认版.docx"
    assert result.is_file()
    assert (project_dir / "08_输出" / "旁白确认版.docx").read_bytes() == result.read_bytes()
    document = Document(result)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    positions = [text.index(f"第{order}页主题") for order in range(1, 9)]
    assert positions == sorted(positions)
    assert "缩略图不可用" in text
    assert len(document.inline_shapes) == 7
    with ZipFile(result) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
        styles = archive.read("word/styles.xml").decode("utf-8")
    assert "Noto Sans SC Thin" in xml + styles


def test_repeated_export_overwrites_fixed_paths_without_numbered_duplicates(
    tmp_path: Path,
) -> None:
    project_dir = tmp_path / "重复导出_20260803_1200"
    manifest = _confirmed_project(project_dir)

    first = export_narration_docx(manifest, project_dir)
    second = export_narration_docx(manifest, project_dir)

    assert first == second
    assert sorted(path.name for path in (project_dir / "04_旁白").glob("*.docx")) == [
        "旁白确认版.docx"
    ]
    assert sorted(path.name for path in (project_dir / "08_输出").glob("*.docx")) == [
        "旁白确认版.docx"
    ]


def test_eight_page_document_renders_without_blank_or_truncated_pages(tmp_path: Path) -> None:
    project_dir = tmp_path / "渲染验收_20260803_1200"
    manifest = _confirmed_project(project_dir, long_page=True)
    output = export_narration_docx(manifest, project_dir)
    structure_text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
    assert ("这是用于验证长段落不会截断的旁白。" * 90) in structure_text

    pdf = render_office_to_pdf(output, tmp_path / "rendered")
    rendered = fitz.open(pdf)
    try:
        assert len(rendered) >= 9
        page_texts = [page.get_text().strip() for page in rendered]
        assert all(page_texts)
        joined = "".join("\n".join(page_texts).split())
        for order in range(1, 9):
            assert f"第{order}页主题" in joined
        assert joined.count("这是用于验证长段落不会截断的旁白。") >= 89
        assert "版本v1|字数1530|预计时长382.5秒" in joined
    finally:
        rendered.close()
