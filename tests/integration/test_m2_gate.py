from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient
from PIL import Image
from workbench.main import create_app
from workbench.ocr.paddle_adapter import OcrResult
from workbench.parsers.pdf_parser import OcrPolicy
from workbench.services.import_service import ImportService
from workbench.services.material_processing_service import MaterialProcessingService
from workbench.services.project_service import ProjectService


class SequencedOcr:
    def __init__(self) -> None:
        self.calls = 0

    def recognize(self, _: Image.Image) -> list[OcrResult]:
        values = [("专业概览", 0.96), ("课程体系", 0.79)]
        text, confidence = values[self.calls]
        self.calls += 1
        return [OcrResult(text=text, bbox=(100, 100, 500, 180), confidence=confidence)]


class ForbiddenOcr:
    def recognize(self, _: Image.Image) -> list[OcrResult]:
        raise AssertionError("缓存命中后不应再次调用 OCR")


def docx_bytes() -> bytes:
    document = Document()
    document.add_heading("专业概览", level=1)
    document.add_paragraph("培养目标")
    document.add_heading("课程体系", level=1)
    document.add_paragraph("机器学习")
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def png_bytes(color: str) -> bytes:
    stream = BytesIO()
    Image.new("RGB", (320, 180), color).save(stream, format="PNG")
    return stream.getvalue()


def test_material_pipeline_persists_and_reuses_results_after_restart(tmp_path: Path) -> None:
    projects = ProjectService(tmp_path)
    project = projects.create("M2阶段验收")
    importer = ImportService(projects)
    importer.import_batch(
        project.id,
        [
            ("大纲.docx", docx_bytes()),
            ("第2页.png", png_bytes("blue")),
            ("第1页.png", png_bytes("red")),
        ],
    )
    ocr = SequencedOcr()
    first = MaterialProcessingService(projects, ocr=ocr).process(project.id, OcrPolicy.ALWAYS)

    assert first.cached is False
    assert ocr.calls == 2
    assert [page.order for page in first.pages] == [1, 2]
    assert first.pages[1].needs_confirmation is True
    assert len(first.matches) == 2
    preview_mtimes = [Path(page.preview_path).stat().st_mtime_ns for page in first.pages]
    projects.close()

    reopened = ProjectService(tmp_path)
    second = MaterialProcessingService(reopened, ocr=ForbiddenOcr()).process(
        project.id, OcrPolicy.ALWAYS
    )

    assert second.cached is True
    assert second.cache_key == first.cache_key
    assert [Path(page.preview_path).stat().st_mtime_ns for page in second.pages] == preview_mtimes
    manifest = reopened.get(project.id)
    assert len(manifest.pages) == 2
    assert len(manifest.matches) == 2
    assert manifest.material_cache_key == first.cache_key


def test_material_parse_api_exposes_cache_status_and_persisted_matches(tmp_path: Path) -> None:
    app = create_app(tmp_path, ocr_engine=SequencedOcr())
    with TestClient(app) as client:
        project = client.post("/api/projects", json={"name": "M2 API"}).json()["data"]
        imported = client.post(
            f"/api/projects/{project['id']}/sources",
            files=[
                (
                    "files",
                    (
                        "大纲.docx",
                        docx_bytes(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
                ("files", ("第2页.png", png_bytes("blue"), "image/png")),
                ("files", ("第1页.png", png_bytes("red"), "image/png")),
            ],
        )
        first = client.post(
            f"/api/projects/{project['id']}/materials/parse",
            json={"ocr_policy": "always"},
        )
        second = client.post(
            f"/api/projects/{project['id']}/materials/parse",
            json={"ocr_policy": "always"},
        )

    assert imported.status_code == 200
    assert first.status_code == 200
    assert first.json()["data"]["cached"] is False
    assert len(first.json()["data"]["matches"]) == 2
    assert second.json()["data"]["cached"] is True


def test_material_parse_api_maps_encrypted_pdf_to_structured_blocking_error(
    tmp_path: Path,
) -> None:
    import fitz

    document = fitz.open()
    document.new_page()
    stream = BytesIO()
    document.save(
        stream,
        encryption=fitz.PDF_ENCRYPT_AES_256,
        owner_pw="owner",
        user_pw="reader",
    )
    document.close()
    app = create_app(tmp_path)
    with TestClient(app, raise_server_exceptions=False) as client:
        project = client.post("/api/projects", json={"name": "加密 PDF"}).json()["data"]
        client.post(
            f"/api/projects/{project['id']}/sources",
            files=[
                ("files", ("大纲.docx", docx_bytes(), "application/octet-stream")),
                ("files", ("加密.pdf", stream.getvalue(), "application/pdf")),
            ],
        )
        response = client.post(
            f"/api/projects/{project['id']}/materials/parse",
            json={"ocr_policy": "never"},
        )

    assert response.status_code == 422
    assert response.json()["error"]["blocking"] is True
    assert "加密" in response.json()["error"]["message"]
