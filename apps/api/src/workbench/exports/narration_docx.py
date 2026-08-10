from __future__ import annotations

import os
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any
from uuid import UUID

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from workbench.domain.confirmation import Confirmation
from workbench.domain.models import PageRecord, ProjectManifest

FONT = "Noto Sans SC Thin"
NAVY = RGBColor(31, 58, 95)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(98, 113, 128)


class NarrationExportError(RuntimeError):
    pass


def export_narration_docx(project: ProjectManifest, project_dir: Path) -> Path:
    project_dir = project_dir.resolve()
    confirmations = _validated_confirmations(project)
    document = Document()
    _configure_document(document)
    _add_cover(document, project)
    extractions = {page.id: page for page in project.page_extractions}
    for page in sorted(project.pages, key=lambda item: item.order):
        confirmation = confirmations[page.id]
        extraction = extractions.get(page.id)
        preview = (
            project_dir / extraction.preview_path
            if extraction and extraction.preview_path
            else None
        )
        _add_page_section(document, page, confirmation, preview)

    narration_dir = project_dir / "04_旁白"
    output_dir = project_dir / "08_输出"
    narration_dir.mkdir(parents=True, exist_ok=True)
    output_dir.mkdir(parents=True, exist_ok=True)
    target = narration_dir / "旁白确认版.docx"
    temporary = narration_dir / ".旁白确认版.tmp.docx"
    document.save(str(temporary))
    os.replace(temporary, target)
    output_target = output_dir / target.name
    output_temporary = output_dir / ".旁白确认版.tmp.docx"
    shutil.copyfile(target, output_temporary)
    os.replace(output_temporary, output_target)
    return target


def _validated_confirmations(project: ProjectManifest) -> dict[UUID, Confirmation]:
    current: dict[UUID, Confirmation] = {}
    for page in project.pages:
        narration = page.narration
        if (
            narration is None
            or narration.confirmed_revision_id is None
            or narration.confirmed_revision_id != narration.revision_id
        ):
            raise NarrationExportError(f"第 {page.order} 页旁白尚未确认")
        confirmation = next(
            (
                item
                for item in reversed(project.narration_confirmations)
                if item.page_id == page.id and item.revision_id == narration.revision_id
            ),
            None,
        )
        if confirmation is None:
            raise NarrationExportError(f"第 {page.order} 页缺少不可变确认记录")
        current[page.id] = confirmation
    return current


def _configure_document(document: DocumentObject) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    _set_style(styles["Normal"], 11, RGBColor(28, 40, 52), after=6, line=1.25)
    _set_style(styles["Title"], 30, NAVY, after=8, line=1.0, bold=True)
    _set_style(styles["Subtitle"], 14, MUTED, after=18, line=1.15)
    _set_style(styles["Heading 1"], 16, BLUE, before=18, after=10, line=1.1, bold=True)
    _set_style(styles["Heading 2"], 13, BLUE, before=14, after=7, line=1.1, bold=True)
    meta = styles.add_style("Narration Metadata", WD_STYLE_TYPE.PARAGRAPH)
    _set_style(meta, 9.5, MUTED, after=6, line=1.1)

    header = section.header.paragraphs[0]
    header.text = "PPT VIDEO WORKBENCH  /  旁白确认版"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _format_paragraph_runs(header, 8.5, MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = "本文件由已确认的逐页旁白生成"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _format_paragraph_runs(footer, 8.5, MUTED)


def _set_style(
    style: Any,
    size: float,
    color: RGBColor,
    *,
    before: float = 0,
    after: float = 0,
    line: float = 1.0,
    bold: bool = False,
) -> None:
    font = style.font
    font.name = FONT
    font.size = Pt(size)
    font.color.rgb = color
    font.bold = bold
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    paragraph = style.paragraph_format
    paragraph.space_before = Pt(before)
    paragraph.space_after = Pt(after)
    paragraph.line_spacing = line


def _add_cover(document: DocumentObject, project: ProjectManifest) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(90)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(kicker, "CONFIRMED NARRATION", 10, BLUE, bold=True)
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("旁白确认版")
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(project.name)
    metadata = document.add_paragraph(style="Narration Metadata")
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.add_run(
        f"项目页数：{len(project.pages)}  |  项目更新时间：{_format_time(project.updated_at)}"
    )
    notice = document.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice.paragraph_format.space_before = Pt(110)
    _add_run(notice, "以下内容均绑定明确旁白版本与人工确认记录", 10, MUTED)


def _add_page_section(
    document: DocumentObject,
    page: PageRecord,
    confirmation: Confirmation,
    preview: Path | None,
) -> None:
    heading = document.add_paragraph(style="Heading 1")
    heading.paragraph_format.page_break_before = True
    heading.paragraph_format.keep_with_next = True
    heading_run = heading.add_run(f"第 {page.order} 页  |  {page.title or '未命名页面'}")
    # Keep the branded document style metadata while using a broadly available
    # Windows CJK font for headings.  This avoids LibreOffice extracting glyphs
    # in visual-column order when the thin display font is unavailable.
    heading_run.font.name = "Microsoft YaHei"
    r_fonts = heading_run._element.get_or_add_rPr().rFonts
    if r_fonts is not None:
        r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    if preview is not None and preview.is_file():
        image_paragraph = document.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.paragraph_format.keep_with_next = True
        image_paragraph.add_run().add_picture(str(preview), width=Inches(5.75))
    else:
        _add_thumbnail_placeholder(document)

    narration_heading = document.add_paragraph(style="Heading 2")
    narration_heading.paragraph_format.keep_with_next = True
    narration_heading.add_run("确认旁白")
    narration = page.narration
    if narration is None:
        raise NarrationExportError(f"第 {page.order} 页旁白缺失")
    body = document.add_paragraph(narration.text)
    body.paragraph_format.keep_together = False
    body.paragraph_format.widow_control = True
    _format_paragraph_runs(body, 11, RGBColor(28, 40, 52))
    character_count = len("".join(narration.text.split()))
    estimated = character_count / 4.0
    metadata = document.add_paragraph(style="Narration Metadata")
    metadata.add_run(
        "  |  ".join(
            [
                f"版本 v{narration.version}",
                f"字数 {character_count}",
                f"预计时长 {estimated:.1f} 秒",
                f"确认人 {confirmation.actor}",
                f"确认时间 {_format_time(confirmation.confirmed_at)}",
            ]
        )
    )


def _add_thumbnail_placeholder(document: DocumentObject) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph.paragraph_format.keep_with_next = True
    _shade_paragraph(paragraph, "EEF2F6")
    _add_run(paragraph, "\n缩略图不可用\n系统仍保留本页已确认旁白\n", 11, MUTED, bold=True)


def _shade_paragraph(paragraph: Any, fill: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _format_paragraph_runs(paragraph: Any, size: float, color: RGBColor) -> None:
    for run in paragraph.runs:
        _format_run(run, size, color)


def _add_run(
    paragraph: Any,
    text: str,
    size: float,
    color: RGBColor,
    *,
    bold: bool = False,
) -> None:
    run = paragraph.add_run(text)
    _format_run(run, size, color, bold=bold)


def _format_run(run: Any, size: float, color: RGBColor, *, bold: bool = False) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)


def _format_time(value: datetime) -> str:
    return value.astimezone().strftime("%Y-%m-%d %H:%M")
