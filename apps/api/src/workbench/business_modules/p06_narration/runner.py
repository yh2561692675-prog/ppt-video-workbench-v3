from __future__ import annotations

import argparse
import hashlib
import os
from datetime import UTC, datetime
from pathlib import Path
from typing import Any, Literal, cast
from uuid import UUID, uuid5

from docx import Document
from peripheral_contracts import BusinessResultManifest, ErrorCategory, JobEnvelope

from workbench.business_modules.p06_narration.models import (
    NarrationAssignment,
    NarrationDocxPayload,
    NarrationExportParameters,
    NarrationGenerateParameters,
    NarrationImportParameters,
    NarrationRevisionsPayload,
    ProjectedNarrationRevision,
    assignment_from_draft,
)
from workbench.business_modules.runtime import (
    BusinessExecution,
    BusinessModuleError,
    StagedArtifact,
    business_input_fingerprint,
    business_parameters,
    execute_business_handler,
    project_revision,
)
from workbench.domain.enums import NodeStatus
from workbench.domain.models import AuditEvent, LlmUsageRecord, NarrationRecord, ProjectManifest
from workbench.integrations.llm.client import LlmClient, LlmIntegrationError
from workbench.narration.generator import NarrationGenerationError, NarrationGenerator
from workbench.narration.repository import NarrationRevision

_ENV_BASE_URL = "WORKBENCH_LLM_BASE_URL"
_ENV_API_KEY = "WORKBENCH_LLM_API_KEY"
_ENV_MODEL = "WORKBENCH_LLM_MODEL"
_ENV_PROFILE_ID = "WORKBENCH_LLM_PROFILE_ID"


def normalize_assignments(assignments: list[dict[str, object]]) -> list[dict[str, str]]:
    """Compatibility helper retained for older callers."""
    normalized: list[dict[str, str]] = []
    for item in assignments:
        parsed = NarrationAssignment.model_validate(item)
        normalized.append(
            {"page_id": str(parsed.page_id), "text": parsed.text.strip(), "author": parsed.author}
        )
    return normalized


def safe_parameters(parameters: dict[str, Any]) -> dict[str, Any]:
    secret_keys = ("key", "token", "secret", "authorization", "credential")

    def clean(value: Any) -> Any:
        if isinstance(value, dict):
            return {
                key: (
                    "[REDACTED]"
                    if any(mark in key.lower() for mark in secret_keys)
                    else clean(item)
                )
                for key, item in value.items()
            }
        if isinstance(value, list):
            return [clean(item) for item in value]
        return value

    return cast(dict[str, Any], clean(parameters))


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--request", required=True, type=Path)
    parser.add_argument("--result", required=True, type=Path)
    args = parser.parse_args()
    job = JobEnvelope.model_validate_json(args.request.read_text(encoding="utf-8-sig"))
    execution = execute_business_handler(job, args.result.parent, args.result, "P06", _handle)
    return 0 if execution.outcome == "succeeded" else 1


def _handle(job: JobEnvelope, attempt_root: Path) -> BusinessExecution:
    if job.job_type == "narration.generate":
        return _generate(job)
    if job.job_type == "narration.import":
        return _import(job)
    if job.job_type == "narration.export":
        return _export(job, attempt_root)
    raise ValueError(f"unsupported P06 job type: {job.job_type}")


def _generate(job: JobEnvelope) -> BusinessExecution:
    parameters = NarrationGenerateParameters.model_validate(business_parameters(job))
    profile_id, base_url, api_key, model = _consume_llm_environment(parameters.profile_id)
    try:
        draft = NarrationGenerator(
            LlmClient(), base_url=base_url, api_key=api_key, model=model
        ).generate(parameters.context)
    except NarrationGenerationError as error:
        raise BusinessModuleError(
            str(error), category=ErrorCategory.INPUT, code=error.code.upper(), retryable=False
        ) from error
    except LlmIntegrationError as error:
        category = (
            ErrorCategory.NETWORK
            if error.code in {"llm_timeout", "llm_unavailable"}
            else ErrorCategory.PROVIDER
        )
        raise BusinessModuleError(
            str(error),
            category=category,
            code=error.code.upper(),
            retryable=error.code in {"llm_timeout", "llm_unavailable", "llm_request_failed"},
        ) from error
    public_digest = hashlib.sha256(base_url.encode("utf-8")).hexdigest()[:16]
    return _revision_execution(
        job,
        "generate",
        (assignment_from_draft(parameters, draft),),
        profile_id=profile_id,
        profile_base_url_digest=public_digest,
        profile_model=model,
    )


def _import(job: JobEnvelope) -> BusinessExecution:
    parameters = NarrationImportParameters.model_validate(business_parameters(job))
    return _revision_execution(job, "import", parameters.assignments)


def _revision_execution(
    job: JobEnvelope,
    operation: Literal["generate", "import"],
    assignments: tuple[NarrationAssignment, ...],
    *,
    profile_id: UUID | None = None,
    profile_base_url_digest: str | None = None,
    profile_model: str | None = None,
) -> BusinessExecution:
    fingerprint = business_input_fingerprint(job)
    created_at = job.created_at.astimezone(UTC)
    revisions = tuple(
        _build_revision(item, fingerprint=fingerprint, created_at=created_at)
        for item in assignments
    )
    payload = NarrationRevisionsPayload(
        operation=operation,
        revisions=revisions,
        profile_id=profile_id,
        profile_base_url_digest=profile_base_url_digest,
        profile_model=profile_model,
    )
    result = BusinessResultManifest(
        schema_version="1.0",
        module_id="P06",
        job_type=job.job_type,
        project_id=job.project_id,
        project_revision=project_revision(job),
        input_fingerprint=fingerprint,
        cache_key=hashlib.sha256(f"{fingerprint}:{job.job_type}".encode()).hexdigest(),
        result_type="narration_revisions",
        payload=payload.model_dump(mode="json"),
    )
    return BusinessExecution(result)


def _build_revision(
    assignment: NarrationAssignment, *, fingerprint: str, created_at: datetime
) -> ProjectedNarrationRevision:
    text = assignment.text.strip()
    author = assignment.author.strip()
    if not text or not author:
        raise ValueError("narration text and author are required")
    revision_id = uuid5(assignment.page_id, f"P06:{fingerprint}:{text}")
    character_count = len("".join(text.split()))
    return ProjectedNarrationRevision(
        id=revision_id,
        page_id=assignment.page_id,
        version=assignment.expected_version + 1,
        text=text,
        author=author,
        source_refs=assignment.source_refs,
        insufficiencies=assignment.insufficiencies,
        warnings=assignment.warnings,
        parent_revision_id=assignment.expected_revision_id,
        created_at=created_at,
        character_count=character_count,
        estimated_duration_seconds=round(max(character_count / 4.0, 0.25), 2),
    )


def _consume_llm_environment(expected_profile_id: UUID) -> tuple[UUID, str, str, str]:
    values = {
        name: os.environ.pop(name, "")
        for name in (_ENV_PROFILE_ID, _ENV_BASE_URL, _ENV_API_KEY, _ENV_MODEL)
    }
    try:
        profile_id = UUID(values[_ENV_PROFILE_ID])
    except ValueError as error:
        raise BusinessModuleError(
            "LLM credential environment is unavailable",
            category=ErrorCategory.ENVIRONMENT,
            code="LLM_CREDENTIAL_UNAVAILABLE",
            retryable=False,
        ) from error
    if profile_id != expected_profile_id or not all(
        values[name] for name in (_ENV_BASE_URL, _ENV_API_KEY, _ENV_MODEL)
    ):
        raise BusinessModuleError(
            "LLM credential environment does not match the requested profile",
            category=ErrorCategory.ENVIRONMENT,
            code="LLM_CREDENTIAL_MISMATCH",
            retryable=False,
        )
    return (
        profile_id,
        values[_ENV_BASE_URL],
        values[_ENV_API_KEY],
        values[_ENV_MODEL],
    )


def _export(job: JobEnvelope, attempt_root: Path) -> BusinessExecution:
    parameters = NarrationExportParameters.model_validate(business_parameters(job))
    target = attempt_root / "narration-confirmed.docx"
    document = Document()
    document.add_heading(parameters.project_name, level=0)
    document.add_paragraph("Confirmed narration")
    for page in parameters.pages:
        document.add_heading(f"Page {page.page_order}: {page.page_title or 'Untitled'}", level=1)
        document.add_paragraph(page.text)
        document.add_paragraph(
            f"Revision v{page.version} | Confirmed by {page.confirmed_by} | "
            f"{page.confirmed_at.astimezone().isoformat()}"
        )
    document.save(str(target))
    content = target.read_bytes()
    relative_path = "04_旁白/旁白确认版.docx"
    payload = NarrationDocxPayload(
        operation="export",
        relative_path=relative_path,
        size_bytes=len(content),
        sha256=hashlib.sha256(content).hexdigest(),
        page_count=len(parameters.pages),
    )
    fingerprint = business_input_fingerprint(job)
    result = BusinessResultManifest(
        schema_version="1.0",
        module_id="P06",
        job_type=job.job_type,
        project_id=job.project_id,
        project_revision=project_revision(job),
        input_fingerprint=fingerprint,
        cache_key=hashlib.sha256(f"{fingerprint}:narration_docx".encode()).hexdigest(),
        result_type="narration_docx",
        payload=payload.model_dump(mode="json"),
    )
    return BusinessExecution(result, (StagedArtifact("narration-docx", "docx", target),))


def project_narration_revisions(result: BusinessResultManifest, project_dir: Path) -> None:
    payload = NarrationRevisionsPayload.model_validate(result.payload)
    manifest_path = project_dir / "project.json"
    manifest = ProjectManifest.model_validate_json(manifest_path.read_text(encoding="utf-8"))
    pages = {page.id: page for page in manifest.pages}

    # Validate the entire batch before writing any immutable/current files.
    for revision in payload.revisions:
        page = pages.get(revision.page_id)
        if page is None:
            raise ValueError(f"narration page does not exist: {revision.page_id}")
        current_id = page.narration.revision_id if page.narration else None
        current_version = page.narration.version if page.narration else 0
        if current_id != revision.parent_revision_id or revision.version != current_version + 1:
            raise ValueError("NARRATION_REVISION_CONFLICT: current narration changed")

    for revision in payload.revisions:
        page = pages[revision.page_id]
        stored = NarrationRevision.model_validate(revision.model_dump(mode="python"))
        _write_immutable_revision(project_dir, stored)
        _write_current_revision(project_dir, stored)
        pages[revision.page_id] = page.model_copy(
            update={
                "narration": NarrationRecord(
                    id=stored.id,
                    revision_id=stored.id,
                    text=stored.text,
                    status=NodeStatus.NEEDS_CONFIRMATION,
                    confirmed_revision_id=None,
                    author=stored.author,
                    version=stored.version,
                    source_refs=stored.source_refs,
                    insufficiencies=stored.insufficiencies,
                    warnings=stored.warnings,
                    updated_at=stored.created_at,
                )
            }
        )

    audit_at = max(item.created_at for item in payload.revisions)
    updates: dict[str, Any] = {
        "pages": sorted(pages.values(), key=lambda item: item.order),
        "audit_log": [
            *manifest.audit_log,
            AuditEvent(
                action=f"narration_{payload.operation}_projected",
                occurred_at=audit_at,
                details={
                    "revision_ids": [str(item.id) for item in payload.revisions],
                    "page_ids": [str(item.page_id) for item in payload.revisions],
                    "profile_id": str(payload.profile_id) if payload.profile_id else None,
                },
            ),
        ],
    }
    if payload.operation == "generate":
        if (
            payload.profile_id is None
            or payload.profile_base_url_digest is None
            or payload.profile_model is None
        ):
            raise ValueError("generated narration is missing profile usage metadata")
        updates["llm_usage"] = [
            *manifest.llm_usage,
            LlmUsageRecord(
                profile_id=payload.profile_id,
                base_url_digest=payload.profile_base_url_digest,
                model=payload.profile_model,
                used_at=audit_at,
            ),
        ]
    updated = manifest.model_copy(update=updates)
    temporary = manifest_path.with_name(".project.json.s1.tmp")
    temporary.write_text(updated.model_dump_json(indent=2) + "\n", encoding="utf-8", newline="\n")
    os.replace(temporary, manifest_path)


def _write_immutable_revision(project_dir: Path, revision: NarrationRevision) -> None:
    history = project_dir / "04_旁白" / "历史版本" / str(revision.page_id)
    history.mkdir(parents=True, exist_ok=True)
    target = history / f"{revision.id}.json"
    serialized = revision.model_dump_json(indent=2) + "\n"
    if target.exists():
        existing = NarrationRevision.model_validate_json(target.read_text(encoding="utf-8"))
        if existing != revision:
            raise ValueError("NARRATION_REVISION_COLLISION: immutable revision differs")
        return
    with target.open("x", encoding="utf-8", newline="\n") as handle:
        handle.write(serialized)
        handle.flush()
        os.fsync(handle.fileno())


def _write_current_revision(project_dir: Path, revision: NarrationRevision) -> None:
    current = project_dir / "04_旁白" / "当前版本"
    current.mkdir(parents=True, exist_ok=True)
    target = current / f"{revision.page_id}.json"
    temporary = target.with_suffix(".tmp")
    temporary.write_text(revision.model_dump_json(indent=2) + "\n", encoding="utf-8", newline="\n")
    os.replace(temporary, target)


if __name__ == "__main__":
    raise SystemExit(main())
